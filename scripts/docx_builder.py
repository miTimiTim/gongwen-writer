"""
公文 .docx 生成工具 — gongwen-writer skill 专用
严格按照 references/document-format.md 格式规范生成文档。
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class GongwenDocx:
    """公文文档构建器"""

    def __init__(self):
        self.doc = Document()
        self._setup_page()

    def _setup_page(self):
        """页面设置"""
        section = self.doc.sections[0]
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.7)
        section.right_margin = Cm(2.6)

    def _set_font(self, run, font_name, size_pt, bold=False):
        """设置中文字体"""
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.bold = bold
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        # 插入到最前面，避免被已有设置覆盖
        existing = rPr.find(qn('w:rFonts'))
        if existing is not None:
            rPr.remove(existing)
        rPr.insert(0, rFonts)

    def _set_line_spacing(self, paragraph, pt_val=29):
        """设置固定行距"""
        pPr = paragraph._element.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:line'), str(int(pt_val * 20)))
        spacing.set(qn('w:lineRule'), 'exact')
        # 替换已有 spacing
        existing = pPr.find(qn('w:spacing'))
        if existing is not None:
            pPr.remove(existing)
        pPr.append(spacing)

    def _add_paragraph(self, text, font_name, size_pt, bold=False,
                       alignment=None, indent_first_line=False):
        """通用添加段落"""
        p = self.doc.add_paragraph()
        self._set_line_spacing(p, 29)
        if alignment is not None:
            p.alignment = alignment
        if indent_first_line:
            p.paragraph_format.first_line_indent = Pt(32)
        run = p.add_run(text)
        self._set_font(run, font_name, size_pt, bold)
        return p

    # ── 公开方法 ──

    def add_title(self, text):
        """方案题目：二号(22pt) 方正小标宋 居中"""
        return self._add_paragraph(text, '方正小标宋简体', 22,
                                   alignment=WD_ALIGN_PARAGRAPH.CENTER)

    def add_blank_line(self):
        """空行"""
        return self._add_paragraph('', '仿宋', 16)

    def add_heading_1(self, text):
        """一级标题：三号(16pt) 黑体 加粗，如"一、总体要求" """
        return self._add_paragraph(text, '黑体', 16, bold=True)

    def add_heading_2(self, text):
        """二级标题：三号(16pt) 黑体 加粗，如"（一）背景与意义" """
        return self._add_paragraph(text, '黑体', 16, bold=True)

    def add_heading_3(self, text):
        """三级标题：三号(16pt) 仿宋 加粗，如"1. 建立XX机制" """
        return self._add_paragraph(text, '仿宋', 16, bold=True)

    def add_body(self, text):
        """正文：三号(16pt) 仿宋，首行缩进2字符"""
        return self._add_paragraph(text, '仿宋', 16,
                                   indent_first_line=True)

    def add_table(self, headers, rows):
        """添加表格：表头黑体四号居中，内容仿宋四号"""
        n_rows = len(rows) + 1
        n_cols = len(headers)
        table = self.doc.add_table(rows=n_rows, cols=n_cols, style='Table Grid')

        # 表头
        for i, text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            self._set_line_spacing(p, 24)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            self._set_font(run, '黑体', 14, bold=True)

        # 数据行
        for r, row_data in enumerate(rows):
            for c, text in enumerate(row_data):
                cell = table.rows[r + 1].cells[c]
                cell.text = ''
                p = cell.paragraphs[0]
                self._set_line_spacing(p, 24)
                run = p.add_run(text)
                self._set_font(run, '仿宋', 14)

        return table

    def save(self, path):
        """保存文档"""
        self.doc.save(path)
        print(f'已保存到: {path}')
